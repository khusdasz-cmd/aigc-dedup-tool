#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
l5_iterative.py — 内部迭代降重引擎（skill 通用版）
将降重变为内部自动循环：改写→验证→根据残留加码→再改写→再验证...直到收敛。

用法：
    python3 l5_iterative.py <input.docx> <output.docx>
"""
import sys
import re
import docx
import l4_check


def safe_set_text(para, text):
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def apply_to_doc(doc, func, label=''):
    """对全文每个段落应用 func(text, para_idx) → 新文本或 None"""
    changes = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        new_text = func(text, i)
        if new_text and new_text != text and new_text.strip():
            safe_set_text(para, new_text)
            changes += 1
    if changes > 0 and label:
        print(f'  [{label}] {changes} 处改动')
    return changes


# ============================================================
# Tier 1 — 安全词汇替换（单字/双字词替换）
# ============================================================

def t1_classical_to_colloquial(text, idx):
    """古典→口语 20 对映射"""
    text = re.sub(r'(?<!不)(?<!\s)可(?!能|以|谓|惜|靠|见|想|乐|悲|爱|恨|行|分|信|逆|燃|降|溶|选|视|控|再|持|靠|怕|笑|耻|敬|佩|贵|疑|怕|算|取|得|知|称|谓|见|信|靠|靠)', '可以', text)
    text = re.sub(r'尚不', '还不', text)
    text = re.sub(r'尚未', '还没有', text)
    text = re.sub(r'尚无', '还没有', text)
    text = re.sub(r'均已', '都已经', text)
    text = re.sub(r'均可', '都可以', text)
    text = re.sub(r'均能', '都能', text)
    text = re.sub(r'仍需', '还需要', text)
    text = re.sub(r'按此', '依据这个', text)
    text = re.sub(r'按以下', '按照以下', text)
    text = re.sub(r'称之为', '叫它', text)
    text = re.sub(r'仍未', '仍然没有', text)
    text = re.sub(r'仍较', '仍然比较', text)
    text = re.sub(r'仍具', '仍然具有', text)
    return text


def t1_ai_verbs(text, idx):
    """AI 高频动词替换"""
    replacements = [
        ('构建了', '建立了'), ('构建', '建立'),
        ('揭示了', '说明了'), ('揭示', '说明'),
        ('验证了', '检验了'), ('验证', '检验'),
        ('聚焦于', '关注'), ('聚焦', '关注'),
        ('旨在', '目标是'), ('以期', '从而'),
        ('引入', '加入'), ('纳入', '列入'),
        ('缺乏', '缺少'), ('开展', '进行'),
        ('运用', '使用'), ('证实', '确认'),
        ('呈现出', '表现出'), ('呈现', '表现'),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    return text


def t1_ai_adjectives(text, idx):
    """AI 高频形容词/副词/名词替换"""
    text = re.sub(r'有显著的[。，]', '有明显的作用', text)
    text = re.sub(r'有显著[的，。]', '有明显', text)
    replacements = [
        ('显著的', '明显的'), ('显著地', '明显地'),
        ('极大地', '大大地'), ('极大的', '很大的'),
        ('系统性', '全面'), ('多维度的', '多方面的'),
        ('综合性的', '综合的'), ('不可或缺', '必不可少'),
        ('广阔', '广泛'), ('精准', '准确'),
        ('充分的', '足够的'), ('充分', '足够'),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    return text


# ============================================================
# Tier 2 — 句式拆解（模板/结构替换）
# ============================================================

def t2_cite_verbs(text, idx):
    """引用动词替换：指出/认为/强调 → 更具体的动词"""
    text = re.sub(r'([）\)])\s*,?\s*指出', r'\1做了分析', text)
    text = re.sub(r'([）\)])\s*,?\s*认为', r'\1提出', text)
    text = re.sub(r'([）\)])\s*,?\s*强调', r'\1专门提到', text)
    text = re.sub(r'([）\)])\s*,?\s*发现', r'\1的结论是', text)
    return text


def t2_paragraph_headers(text, idx):
    """段落首尾模板清理"""
    text = re.sub(r'^(先说|再说|再看|再来看|说到)[^。，]{2,15}[，。]', '', text)
    text = re.sub(r'^梳理已有文献可以发现[，,。]', '', text)
    text = re.sub(r'^综合上述分析[，,。]', '', text)
    text = re.sub(r'还需要(强调|提一下|指出)', '还有一个方面值得', text)
    return text


def t2_transition_words(text, idx):
    """替换论文过渡词"""
    text = re.sub(r'首先[，,]', '', text)
    text = re.sub(r'其次[，,]', '另外', text)
    text = re.sub(r'再次[，,]', '再加上', text)
    text = re.sub(r'最后[，,]', '一个关键点是', text)
    return text


def t2_theory_labels(text, idx):
    """去掉理论标签"""
    text = re.sub(r'(从|在)[^。，]{2,12}(理论|视角|维度|层面|角度)(来看|出发|看)[，。]', '', text)
    text = re.sub(r'基于[^。，]{2,15}(理论|视角|框架)[，。]', '', text)
    return text


def t2_packaging(text, idx):
    """去掉包装词"""
    text = re.sub(r'(呈现出|表现出)[^。]{2,15}(特征|趋势|格局|特点)', '', text)
    text = re.sub(r'(具有|有)明显的(正向|促进|推动|提升|抑制|负面|积极)\s*(效应|作用)', '有显著的', text)
    return text


def t2_definition_evolution(text, idx):
    """拆解三阶段单句"""
    text = re.sub(
        r'(概念|内涵)(经历|经历)了(从[^。]{2,15}到[^。]{2,15}再到[^。]{2,15})的(演变|发展|过程)(?!\w)',
        r'\1的理解在不断拓宽。\3的变化反映了认识的深化',
        text
    )
    text = re.sub(
        r'(这个|该)概念经历了(从[^。]{2,15}到[^。]{2,15}再到[^。]{2,15})的演变(过程)?',
        r'人们对这个概念的理解在逐步拓宽。\2的变化反映了不同阶段的理论关注点',
        text
    )
    return text


def t2_method_definition(text, idx):
    """拆解方法论定义"""
    text = re.sub(r'是一种[^。]{3,25}的方法，其基本思想是', r'的核心思路是', text)
    text = re.sub(r'其基本思想是根据[^。]{5,40}来确定权重', r'基本逻辑是：指标数值差异越大，区分度越高，权重就越大', text)
    return text


def t2_column_labels(text, idx):
    """列(X) → 第X个模型"""
    text = re.sub(r'列（(\d)）[的,，]*(结果|显示|报告)', r'第\1个模型的结果', text)
    text = re.sub(r'列\((\d)\)[的,，]*(结果|显示|报告)', r'第\1个模型的结果', text)
    return text


def t2_purpose_sentences(text, idx):
    """目的句替换"""
    text = re.sub(r'^为了保证[^。]{3,20}可靠[，。]', '', text)
    text = re.sub(r'^为了进一步缓解[^。]{3,30}[，。]', '', text)
    text = re.sub(r'^为了检验[^。]{5,30}[，。]', '', text)
    return text


def t2_statistical_quadruplet(text, idx):
    """打断统计量四连"""
    text = re.sub(
        r'(标准差为[^，。]+)，最小值为[^，。]+，最大值为[^。]+',
        r'\1；最小和最大值也列在了表里',
        text
    )
    return text


def t2_ratio_template(text, idx):
    """占比模板自然化"""
    text = re.sub(r'占比达到(\d+\.?\d*)%', r'\1%的部分来自', text)
    return text


# ============================================================
# Tier 3 — 段落级改写（结构清零）
# ============================================================

def t3_bureaucratic(text, idx):
    """去公文腔"""
    text = re.sub(r'持续推进', '一直在推', text)
    text = re.sub(r'系统推进', '整体推进', text)
    text = re.sub(r'深入推进', '进一步推进', text)
    text = re.sub(r'突出抓好', '重点管好', text)
    text = re.sub(r'(筑牢|守住).{0,4}(安全|底线)', '守好安全', text)
    text = re.sub(r'健康可持续发展', '持续发展', text)
    text = re.sub(r'发挥着重要的引领和示范作用', '走在前列', text)
    text = re.sub(r'提供了坚实的资金保障', '提供了资金支持', text)
    text = re.sub(r'牢牢把握', '抓住', text)
    text = re.sub(r'着力打造', '努力建设', text)
    return text


def t3_numbered_steps(text, idx):
    """步骤编号清零"""
    text = re.sub(r'^第一步[，,]\s*', '先做', text)
    text = re.sub(r'^第二步[，,]\s*', '再', text)
    text = re.sub(r'^第三步[，,]\s*', '然后', text)
    text = re.sub(r'^第四步[，,]\s*', '下一步', text)
    text = re.sub(r'^第五步[，,]\s*', '接着', text)
    text = re.sub(r'^第六步[，,]\s*', '最后', text)
    return text


def t3_policy_suggestions(text, idx):
    """政策建议去'应/要'"""
    text = re.sub(r'应加快', '加快', text)
    text = re.sub(r'应进一步', '进一步', text)
    text = re.sub(r'应加强', '加强', text)
    text = re.sub(r'应完善', '完善', text)
    text = re.sub(r'应推动', '推动', text)
    text = re.sub(r'应继续', '继续', text)
    text = re.sub(r'应加大', '加大', text)
    text = re.sub(r'应深化', '深化', text)
    text = re.sub(r'应该', '', text)
    return text


# ============================================================
# Adaptive — 靶向修复（根据残留问题类型激活）
# ============================================================

def adaptive_numeral(text, idx):
    """编号列举靶向修复"""
    text = re.sub(r'^第一[，,]\s*', '', text)
    text = re.sub(r'^第二[，,]\s*', '再看', text)
    text = re.sub(r'^第三[，,]\s*', '最后', text)
    text = re.sub(r'^第四[，,]\s*', '此外', text)
    text = re.sub(r'^第五[，,]\s*', '还有一个方面', text)
    return text


def adaptive_compression(text, idx):
    """介词压缩结构拆解"""
    text = re.sub(r'对([^。，]{2,15})进行([^。，]{2,6})', r'\2\1', text)
    if len(text) > 20:
        text = re.sub(r'在([^。，]{2,10})方面', r'在\1上', text)
    return text


def adaptive_dense_citation(text, idx):
    """密集引用压缩：一句内超过2个引用→保留前两个"""
    cites = list(re.finditer(r'[（(][^）)]{2,20}[，,]\d{4}[）)]', text))
    if len(cites) >= 3:
        for m in reversed(cites[2:]):
            before = text[:m.start()]
            after = text[m.end():]
            if after.startswith('）') or after.startswith(')'):
                after = after[1:]
            text = before + after
    return text


def adaptive_yishisan(text, idx):
    """'一是/二是/三是' 枚举替换"""
    text = re.sub(r'([，。])一是([^，。]{3,20})，', r'\1第一个方面是\2，', text)
    text = re.sub(r'([，。])二是([^，。]{3,20})[，。]', r'\1另外，\2', text)
    text = re.sub(r'([，。])三是([^，。]{3,20})', r'\1还有一个是\2', text)
    return text


def adaptive_article_ref(text, idx):
    """段落开头'本文/本研究'变换"""
    if idx % 3 == 0:
        text = re.sub(r'^(本文|本研究)', '这篇论文', text)
    elif idx % 3 == 1:
        text = re.sub(r'^(本文|本研究)', '这次研究', text)
    return text


def adaptive_formal_opening(text, idx):
    """正式开场白削减"""
    text = re.sub(r'从([^。，]{2,8})的角度来看', r'从\1的角度', text)
    text = re.sub(r'在([^。，]{2,10})层面', r'在\1方面', text)
    return text


# ============================================================
# 模式注册表
# ============================================================

TIER1_PATTERNS = [
    ('古典→口语', t1_classical_to_colloquial),
    ('AI动词替换', t1_ai_verbs),
    ('AI形/名替换', t1_ai_adjectives),
]

TIER2_PATTERNS = [
    ('引用动词', t2_cite_verbs),
    ('段落首尾', t2_paragraph_headers),
    ('过渡词', t2_transition_words),
    ('理论标签', t2_theory_labels),
    ('包装词', t2_packaging),
    ('定义演变', t2_definition_evolution),
    ('方法论定义', t2_method_definition),
    ('列标签', t2_column_labels),
    ('目的句', t2_purpose_sentences),
    ('统计量四连', t2_statistical_quadruplet),
    ('占比模板', t2_ratio_template),
]

TIER3_PATTERNS = [
    ('去公文腔', t3_bureaucratic),
    ('步骤编号', t3_numbered_steps),
    ('政策建议', t3_policy_suggestions),
]

ADAPTIVE_PATTERNS = [
    ('编号清零', adaptive_numeral, '编号列举', 1),
    ('介词拆解', adaptive_compression, '介词压缩结构', 3),
    ('密集引用压缩', adaptive_dense_citation, '密集引用', 2),
    ('一是二是三是', adaptive_yishisan, '枚举结构', 1),
    ('本文/本研究', adaptive_article_ref, '自指开头', 3),
    ('正式开场白', adaptive_formal_opening, '正式开场白', 2),
]


# ============================================================
# 主迭代循环
# ============================================================

MAX_ITERATIONS = 15
CONVERGENCE_WINDOW = 3
IMPROVEMENT_THRESHOLD = 0.03
MAX_TIER_PASSES = 4


def apply_patterns(doc, patterns, label=''):
    """应用一组 pattern 到全文"""
    changes = 0
    for name, func in patterns:
        c = apply_to_doc(doc, func, label=f'{label}/{name}')
        changes += c
    return changes


def run_all_tiers(doc):
    """依次应用 T1→T2→T3，每级反复直到不再变化"""
    total_changes = 0
    for tier_num, (patterns, label) in enumerate(
        [(TIER1_PATTERNS, 'T1'), (TIER2_PATTERNS, 'T2'), (TIER3_PATTERNS, 'T3')], 1
    ):
        for passage in range(MAX_TIER_PASSES):
            c = apply_patterns(doc, patterns, label)
            total_changes += c
            if c == 0:
                break
    return total_changes


def run_adaptive(doc, previous_issues):
    """根据残留问题靶向修复"""
    changes = 0
    if not previous_issues:
        return 0
    by_cat = previous_issues['by_category']
    for name, func, req_cat, min_count in ADAPTIVE_PATTERNS:
        if req_cat in by_cat and by_cat[req_cat]['count'] >= min_count:
            c = apply_to_doc(doc, func, label=f'ADA/{name}')
            changes += c
    return changes


def is_converged(history):
    """检查是否收敛"""
    if len(history) < CONVERGENCE_WINDOW:
        return False
    for h in history[-3:]:
        if h['changes'] == 0:
            return True
    if history[-1]['total'] == 0:
        return True
    recent = history[-CONVERGENCE_WINDOW:]
    first_w = recent[0]['weighted']
    last_w = recent[-1]['weighted']
    if first_w > 0:
        improvement = (first_w - last_w) / first_w
        if improvement < IMPROVEMENT_THRESHOLD:
            return True
    return False


# ============================================================
# 主入口
# ============================================================

def main():
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    print(f"{'='*60}")
    print(f"l5_iterative.py — 内部迭代降重引擎 (skill 通用版)")
    print(f"{'='*60}")
    print(f"输入: {input_path}")
    print(f"输出: {output_path}")
    print(f"{'='*60}\n")

    doc = docx.Document(input_path)

    # 初始检查
    result = l4_check.check_all(doc)
    history = [{
        'iteration': 0,
        'tier': 0,
        'changes': 0,
        'total': result['total'],
        'weighted': result['weighted_total'],
        'categories': {k: v['count'] for k, v in result['by_category'].items()},
    }]

    print(f"初始状态:")
    print(f"  原始问题数: {result['total']}")
    print(f"  加权总分:   {result['weighted_total']:.1f}\n")

    print(f"第1阶段: 应用所有改写模式 (T1→T2→T3)...")
    phase1_changes = run_all_tiers(doc)
    result = l4_check.check_all(doc)
    print(f"  T1→T2→T3 共 {phase1_changes} 处改动")
    print(f"  剩余: 原始{result['total']}处, 加权{result['weighted_total']:.1f}\n")

    history.append({
        'iteration': 1, 'tier': 'ALL', 'changes': phase1_changes,
        'total': result['total'], 'weighted': result['weighted_total'],
        'categories': {k: v['count'] for k, v in result['by_category'].items()},
    })

    print(f"第2阶段: adaptive 靶向迭代...")
    for iteration in range(2, MAX_ITERATIONS + 1):
        prev_issues = result
        prev_total = result['total']
        changes = run_adaptive(doc, prev_issues)
        if changes == 0:
            print(f"  第{iteration}轮: 无需进一步修复，停止")
            break

        result = l4_check.check_all(doc)
        is_improving = result['total'] < prev_total

        history.append({
            'iteration': iteration, 'tier': 'ADA', 'changes': changes,
            'total': result['total'], 'weighted': result['weighted_total'],
            'categories': {k: v['count'] for k, v in result['by_category'].items()},
        })

        delta = result['total'] - prev_total
        arrow = '↓' if is_improving else '↑'
        print(f"  第{iteration}轮 [ADA] | 改动{changes:3d}处 | "
              f"原始{result['total']:3d}({arrow}{delta:+d}) | "
              f"加权{result['weighted_total']:6.1f}")

        if is_converged(history):
            print(f"  ✓ 已收敛")
            break

        extra = apply_patterns(doc, TIER2_PATTERNS, 'T2-retry')
        extra += apply_patterns(doc, TIER3_PATTERNS, 'T3-retry')
        if extra > 0:
            result = l4_check.check_all(doc)
            print(f"    回溯应用 T2/T3: +{extra} 处")
        else:
            break

    final = history[-1]
    initial = history[0]

    print(f"\n{'='*60}")
    print(f"迭代完成 — 共 {final['iteration']} 轮")
    print(f"{'='*60}")
    print(f"{'':20s} {'初始':>8s} {'最终':>8s} {'变化':>8s}")
    print(f"{'─'*60}")
    print(f"{'原始问题数':20s} {initial['total']:8d} {final['total']:8d} {final['total']-initial['total']:+7d}")
    print(f"{'加权总分':20s} {initial['weighted']:8.1f} {final['weighted']:8.1f} {final['weighted']-initial['weighted']:+7.1f}")
    print(f"{'─'*60}")
    for cat_name in sorted(final['categories'].keys()):
        ic = initial['categories'].get(cat_name, 0)
        fc = final['categories'].get(cat_name, 0)
        if ic > 0 or fc > 0:
            print(f"{cat_name:20s} {ic:8d} {fc:8d} {fc-ic:+7d}")
    print(f"{'─'*60}")

    doc.save(output_path)
    print(f"\n已保存: {output_path}")
    print(f"{'='*60}")

    if final['total'] > 0:
        print(f"\n⚠ 残留 {final['total']} 处，建议检查：")
        all_issues = l4_check.check_all_by_para(doc)
        for para_idx in sorted(all_issues.keys())[:10]:
            cats = set(cat for cat, _, _ in all_issues[para_idx])
            print(f"  段落 [{para_idx}]: {', '.join(cats)}")
        if len(all_issues) > 10:
            print(f"  ... 共 {len(all_issues)} 个段落有残留")


if __name__ == '__main__':
    main()
