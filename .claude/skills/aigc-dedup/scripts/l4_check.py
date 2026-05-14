#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
l4_check.py — AIGC 降重检查引擎（可导入模块）
从 l4_verify.py 提取，返回结构化检查结果供 l5_iterative.py 使用
"""
import re
import docx

# ======== Check 1: 文言单字 ========
CLASSICAL_PATTERNS = [
    (r'(?<![一-鿿])可[一-鿿　-〿＀-￯]{2,}', '可→可以（非术语语境）'),
    ('尚不', '尚→还'), ('尚未', '尚→还'), ('尚无', '尚→还'),
    ('均已', '均→都'), ('均可', '均→都'), ('均能', '均→都'),
    ('均呈现', '均→都'), ('均需要', '均→都'),
    ('按此', '按→依据'), ('按以下', '按→按照'),
    ('称之为', '称→叫'), ('正如', '如→比如'),
    ('仍未', '仍→仍然'), ('仍较', '仍→仍然'), ('仍具', '仍→仍然'),
]

def check_classical(doc):
    issues = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        for pattern, label in CLASSICAL_PATTERNS:
            if re.search(pattern, text):
                match = re.search(pattern, text)
                start = max(0, match.start() - 20)
                end = min(len(text), match.end() + 30)
                ctx = text[start:end]
                issues.append((i, label, ctx))
                break
    return issues

# ======== Check 2: AI高频动词 ========
HIGH_FREQ_VERBS = [
    ('构建', '→ 建立/搭建'), ('揭示', '→ 说明/显示'),
    ('验证', '→ 检验/证实'), ('运用', '→ 使用'),
    ('构建了', '→ 建立了'), ('旨在', '→ 目的是'),
    ('以期', '→ 希望/从而'), ('引入', '→ 加入/纳入'),
    ('纳入', '→ 列入'), ('缺乏', '→ 缺少'),
    ('开展', '→ 进行/做'), ('聚焦', '→ 关注/围绕'),
]

def check_verbs(doc):
    issues = []
    seen = set()
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        for word, label in HIGH_FREQ_VERBS:
            if word in text and (i, word) not in seen:
                seen.add((i, word))
                match = re.search(word, text)
                start = max(0, match.start() - 15)
                end = min(len(text), match.end() + 25)
                ctx = text[start:end]
                issues.append((i, f'「{word}」{label}', ctx))
    return issues

# ======== Check 3: AI高频名词/形容词 ========
HIGH_FREQ_NOUNS = [
    ('显著的', '→ 明显的'), ('极大地', '→ 大大地'),
    ('系统性', '→ 全面'), ('多维度的', '→ 多个方面的'),
    ('综合性的', '→ 综合的'), ('不可或缺', '→ 必不可少'),
    ('关键的', '→ 重要的/核心的'), ('策略', '→ 方案/办法'),
    ('机制', '→ 原理/方式'), ('思路', '→ 方法/方案'),
    ('视角', '→ 角度'), ('维度', '→ 方面/角度'),
    ('广泛关注', '→ 很多关注'), ('精准', '→ 准确'), ('广阔', '→ 广泛'),
]

def check_nouns(doc):
    issues = []
    seen = set()
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        for word, label in HIGH_FREQ_NOUNS:
            if word in text and (i, word) not in seen:
                seen.add((i, word))
                match = re.search(word, text)
                start = max(0, match.start() - 15)
                end = min(len(text), match.end() + 25)
                ctx = text[start:end]
                issues.append((i, f'「{word}」{label}', ctx))
    return issues

# ======== Check 4: 套话句式 ========
TEMPLATE_PATTERNS = [
    (r'不仅[^。]*而且', '不仅…而且… 未拆解'),
    (r'随着[^。]{5,30}的不断(发展|推进|深化)', '随着…的不断发展'),
    (r'扮演[^。]*角色', '扮演着重要的角色'),
    (r'引起了广泛关注', '引起了广泛关注'),
    (r'具有重要[^。]*(理论意义|现实意义)', '具重要理论/现实意义'),
    (r'在当今[^。]*背景下', '在当今…背景下'),
    (r'为[^。]{2,20}提供了新的思路', '提供了新的思路'),
    (r'[^。]*重要途径之一', '是…的重要途径之一'),
    (r'打下了良好的基础', '打下了良好的基础'),
    (r'整体规律十分清晰', '整体规律十分清晰'),
    (r'综上所述', '综上所述'),
    (r'为[^。]*注入了新的动力', '注入新动力'),
    (r'具有决定性影响', '具决定性影响'),
    (r'发挥了不可替代的作用', '不可替代的作用'),
]

def check_templates(doc):
    issues = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        for pattern, label in TEMPLATE_PATTERNS:
            if re.search(pattern, text):
                match = re.search(pattern, text)
                start = max(0, match.start() - 10)
                end = min(len(text), match.end() + 20)
                ctx = text[start:end]
                issues.append((i, label, ctx))
    return issues

# ======== Check 5: 编号列举 ========
NUMERAL_PATTERNS = [
    (r'(第一|第二|第三|第四|第五)[，,、]', '编号列举'),
    (r'（(1|2|3|4|5|6)）[^）]{5,30}显著', '编号的显著描述'),
    (r'列（\d）为', '列(1)/列(2) 描述'),
    (r'^第[一二三四五六]步[，,]', '几步几步骤（段落开头）'),
]

def check_numeral(doc):
    issues = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        for pattern, label in NUMERAL_PATTERNS:
            if re.search(pattern, text):
                match = re.search(pattern, text)
                start = max(0, match.start() - 10)
                end = min(len(text), match.end() + 30)
                ctx = text[start:end]
                issues.append((i, label, ctx))
    return issues

# ======== Check 6: 控制变量排比 ========
def check_parallel(doc):
    issues = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        pos_count = len(re.findall(r'显著为正', text))
        neg_count = len(re.findall(r'显著为负', text))
        if pos_count + neg_count >= 3:
            issues.append((i, f'含{pos_count+neg_count}处显著为正/负排比', text[:100]))
    return issues

# ======== Check 7: 介词压缩 ========
COMPRESSION_PATTERNS = [
    (r'其[一-鿿]{2,6}[的]', '其+名 需拆解'),
    (r'以[一-鿿]{2,12}为[一-鿿]', '以A为B 需拆解'),
    (r'对[一-鿿，]{2,15}进行[一-鿿]{2,4}', '对X进行Y 需简化'),
    (r'在[一-鿿，]{2,15}方面', '在X方面→在X上'),
]

def check_compression(doc):
    issues = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        for pattern, label in COMPRESSION_PATTERNS:
            if re.search(pattern, text):
                match = re.search(pattern, text)
                start = max(0, match.start() - 10)
                end = min(len(text), match.end() + 20)
                ctx = text[start:end]
                issues.append((i, label, ctx))
    return issues

# ======== Check 8: 三阶段演变 ========
STAGE_3_RE = re.compile(r'从[^。]{2,15}到[^。]{2,15}再到[^。]{2,15}')
STAGE_HISTORY = re.compile(r'经历[^。]{3,20}阶段')

def check_stage3(doc):
    issues = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if STAGE_3_RE.search(text) or STAGE_HISTORY.search(text):
            match = STAGE_3_RE.search(text) or STAGE_HISTORY.search(text)
            start = max(0, match.start() - 10)
            end = min(len(text), match.end() + 30)
            ctx = text[start:end]
            issues.append((i, '三阶段演变', ctx))
    return issues

# ======== Check 9: 方法论定义 ========
METHOD_PATTERNS = [
    (r'是一种[^。]{3,30}的方法，其基本思想是', '是一种...的方法，其基本思想是'),
    (r'其基本思想是[^。]*指标[^。]*信息熵[^。]*越小', '其基本思想(熵权法)'),
]

def check_method(doc):
    issues = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        for pattern, label in METHOD_PATTERNS:
            if re.search(pattern, text):
                match = re.search(pattern, text)
                start = max(0, match.start() - 10)
                end = min(len(text), match.end() + 30)
                ctx = text[start:end]
                issues.append((i, label, ctx))
    return issues

# ======== Check 10: 统计数字密集 ========
def check_stats(doc):
    issues = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        stats_count = len(re.findall(r'(均值|标准差|最小值|最大值|中位数|方差)\s*(为|=|:|：)', text))
        if stats_count >= 3:
            issues.append((i, f'密集统计量({stats_count}个)', text[:120]))
        col_count = len(re.findall(r'[列（(][1-5][）)]', text))
        if col_count >= 3:
            issues.append((i, f'列编号({col_count}处)', text[:120]))
    return issues

# ======== Check 11: 引用动词 ========
CITE_VERB_RE = re.compile(r'[））\)][，,]\s*(指出|认为|强调)')

def check_cite_verbs(doc):
    issues = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        matches = CITE_VERB_RE.findall(text)
        for m in matches:
            issues.append((i, f'引用后跟「{m}」', text[:120]))
    return issues


# ======== Check 12: 密集引用 ========
def check_dense_citation(doc):
    issues = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        # 检测一句内有3+个(Author, Year)引用
        cites = re.findall(r'[（(][^）)]{2,20}[，,]\d{4}[）)]', text)
        if len(cites) >= 3:
            match = re.search(r'[（(][^）)]{2,20}[，,]\d{4}[）)]', text)
            start = max(0, match.start() - 10)
            end = min(len(text), match.end() + 30)
            ctx = text[start:end]
            issues.append((i, f'密集引用({len(cites)}处)', ctx))
    return issues

# ======== Check 13: 枚举结构(一是/二是/三是) ========
def check_yishisan(doc):
    issues = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if re.search(r'一是[^。，]{3,30}[，。]二是', text):
            issues.append((i, '一是二是三是枚举', text[:120]))
        elif re.search(r'头一条[。，]', text) or re.search(r'是第[一二三]条', text):
            issues.append((i, '头一条/第二条枚举', text[:120]))
    return issues

# ======== Check 14: 自指开头(本文/本研究) ========
def check_self_ref(doc):
    issues = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if re.match(r'^(本文|本研究)', text):
            issues.append((i, '自指开头', text[:60]))
    return issues

# ======== Check 15: 过渡词密集 ========
def check_transition_density(doc):
    issues = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text or len(text) < 50:
            continue
        t_count = len(re.findall(r'[而也则]', text))
        if t_count >= 5:
            issues.append((i, f'过渡词密集({t_count}处)', text[:100]))
    return issues

# ======== Check 16: 正式开场白 ========
def check_formal_opening(doc):
    issues = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if re.search(r'从[^。，]{2,8}的角度来看', text):
            issues.append((i, '从X角度来看开场白', text[:80]))
        if re.search(r'在[^。，]{2,10}层面', text):
            issues.append((i, '在X层面句式', text[:80]))
    return issues


# ======== 汇总 ========

CHECK_FUNCTIONS = [
    ('文言单字', check_classical),
    ('AI高频动词', check_verbs),
    ('AI高频名词', check_nouns),
    ('套话句式', check_templates),
    ('编号列举', check_numeral),
    ('控制变量排比', check_parallel),
    ('介词压缩结构', check_compression),
    ('三阶段演变', check_stage3),
    ('方法论定义', check_method),
    ('统计数字密集', check_stats),
    ('引用动词', check_cite_verbs),
    ('密集引用', check_dense_citation),
    ('枚举结构', check_yishisan),
    ('自指开头', check_self_ref),
    ('过渡词密集', check_transition_density),
    ('正式开场白', check_formal_opening),
]

# 各类别权重（用于计算加权总分，高/中/低影响 PaperPass 分数）
CATEGORY_WEIGHTS = {
    '编号列举': 3.0,
    '统计数字密集': 3.0,
    '套话句式': 3.0,
    '方法论定义': 3.0,
    '三阶段演变': 2.0,
    '文言单字': 1.5,
    'AI高频动词': 1.5,
    '引用动词': 1.5,
    '控制变量排比': 1.5,
    '密集引用': 1.5,
    '枚举结构': 1.5,
    '自指开头': 1.0,
    '过渡词密集': 1.0,
    '正式开场白': 1.0,
    'AI高频名词': 1.0,
    '介词压缩结构': 1.0,
}


def check_all(doc):
    """
    对所有检查项执行检查，返回结构化结果：
    {
        'by_category': { '类别名': { 'count': int, 'issues': [(idx, label, ctx), ...] } },
        'weighted_total': float,  # 加权总分
        'total': int,             # 原始总分
    }
    """
    result = {'by_category': {}, 'weighted_total': 0.0, 'total': 0}

    for name, func in CHECK_FUNCTIONS:
        issues = func(doc)
        count = len(issues)
        weight = CATEGORY_WEIGHTS.get(name, 1.0)
        result['by_category'][name] = {
            'count': count,
            'issues': issues,
            'weight': weight,
        }
        result['weighted_total'] += count * weight
        result['total'] += count

    return result


def check_all_by_para(doc):
    """
    按段落汇总问题，返回 { 段落索引: [(类别, 详情, 上下文), ...] }
    """
    para_issues = {}
    for name, func in CHECK_FUNCTIONS:
        issues = func(doc)
        for idx, label, ctx in issues:
            if idx not in para_issues:
                para_issues[idx] = []
            para_issues[idx].append((name, label, ctx))
    return para_issues


if __name__ == '__main__':
    import sys
    input_path = sys.argv[1] if len(sys.argv) > 1 else r'D:\desktop\扩充c1\output\毕业论文7_L3降重v9.docx'
    doc = docx.Document(input_path)
    result = check_all(doc)

    print(f"{'='*60}")
    print(f"AIGC 降重检查报告")
    print(f"{'='*60}")
    print(f"文件: {input_path}")
    print(f"{'='*60}\n")

    for name, data in result['by_category'].items():
        count = data['count']
        print(f"{name:12s}: {count:3d} 处", end='')
        if count > 0:
            idxs = sorted(set(idx for idx, _, _ in data['issues']))
            print(f"  (段落: {idxs[:5]}{'...' if len(idxs) > 5 else ''})", end='')
        print()

    print(f"{'─'*60}")
    print(f"原始总计: {result['total']} 处")
    print(f"加权总计: {result['weighted_total']:.1f}")
    print(f"{'='*60}")
